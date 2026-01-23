"""
Computer Controller Platform Extension

Provides computer control tools:
- web_scrape: Fetch and save web content
- automation_script: Run automation scripts
- computer_control: System automation
- xlsx_tool: Process Excel files
- docx_tool: Process Word documents
- pdf_tool: Process PDF files
- cache: Manage cached files

Reference: goose-rs/crates/goose-mcp/src/computercontroller/mod.rs
"""

import asyncio
import hashlib
import json
import os
import platform as sys_platform
import shutil
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional
from urllib.parse import urlparse

try:
    import httpx
    HTTPX_AVAILABLE = True
except ImportError:
    HTTPX_AVAILABLE = False


class ComputerControllerPlatformExtension:
    """Computer Controller Platform Extension"""

    EXTENSION_NAME = "computer_controller"

    def __init__(self):
        self._initialized = False
        self._os = os.name
        self.cache_dir = self._get_cache_dir()
        self.http_client = None

    def _get_cache_dir(self) -> Path:
        """Get cache directory"""
        if self._os == "nt":
            base = Path(os.environ.get("LOCALAPPDATA", ""))
            return base / "Block" / "goose" / "cache" / "computer_controller"
        else:
            cache_home = os.environ.get("XDG_CACHE_HOME", "")
            if cache_home:
                return Path(cache_home) / "goose" / "computer_controller"
            return Path.home() / ".cache" / "goose" / "computer_controller"

    def _ensure_cache_dir(self) -> None:
        """Ensure cache directory exists"""
        self.cache_dir.mkdir(parents=True, exist_ok=True)

    def _get_cache_path(self, prefix: str, extension: str) -> Path:
        """Generate cache file path"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return self.cache_dir / f"{prefix}_{timestamp}.{extension}"

    async def initialize(self) -> Dict[str, Any]:
        """Initialize the extension"""
        self._ensure_cache_dir()

        if HTTPX_AVAILABLE:
            self.http_client = httpx.AsyncClient(timeout=30.0, follow_redirects=True)

        instructions = self._generate_instructions()
        self._initialized = True

        return {
            "name": self.EXTENSION_NAME,
            "version": "1.0.0",
            "description": "Computer control and automation tools",
            "instructions": instructions,
        }

    def _generate_instructions(self) -> str:
        """Generate instructions based on OS"""
        os_specific = ""
        if self._os == "nt":
            os_specific = """Extra tools:
- automation_script: Create and run PowerShell or Batch scripts
- computer_control: System automation using PowerShell"""
        elif sys_platform.system() == "Darwin":
            os_specific = """Extra tools:
- automation_script: Create and run Shell and Ruby scripts
- computer_control: System automation using AppleScript"""
        else:
            os_specific = """Extra tools:
- automation_script: Create and run Shell scripts
- computer_control: System automation using shell commands"""

        return f"""You are a helpful assistant for common tasks like web scraping,
data processing, and automation.

You can use scripting to work with text files, CSVs, JSON, etc.
Accessing websites and APIs is common.

{os_specific}

Cache directory: {self.cache_dir}"""

    async def list_tools(self) -> List[Dict[str, Any]]:
        """List available tools"""
        if not self._initialized:
            await self.initialize()

        tools = [
            {
                "name": "web_scrape",
                "description": "Fetch and save content from a web page (text, JSON, or binary)",
                "inputSchema": {
                    "type": "object",
                    "properties": {
                        "url": {"type": "string", "description": "URL to fetch"},
                        "save_as": {"type": "string", "enum": ["text", "json", "binary"], "default": "text"},
                    },
                    "required": ["url"],
                }
            },
            {
                "name": "cache",
                "description": "Manage cached files: list, view, delete, clear",
                "inputSchema": {
                    "type": "object",
                    "properties": {
                        "command": {"type": "string", "enum": ["list", "view", "delete", "clear"]},
                        "path": {"type": "string", "description": "Path for view/delete"},
                    },
                    "required": ["command"],
                }
            },
        ]

        if self._os == "nt":
            tools.append({
                "name": "automation_script",
                "description": "Create and run PowerShell or Batch scripts",
                "inputSchema": {
                    "type": "object",
                    "properties": {
                        "language": {"type": "string", "enum": ["powershell", "batch"]},
                        "script": {"type": "string", "description": "Script content"},
                        "save_output": {"type": "boolean", "default": False},
                    },
                    "required": ["language", "script"],
                }
            })
            tools.append({
                "name": "computer_control",
                "description": "Control computer using PowerShell automation",
                "inputSchema": {
                    "type": "object",
                    "properties": {
                        "script": {"type": "string", "description": "PowerShell script"},
                        "save_output": {"type": "boolean", "default": False},
                    },
                    "required": ["script"],
                }
            })
        else:
            tools.append({
                "name": "automation_script",
                "description": "Create and run Shell scripts",
                "inputSchema": {
                    "type": "object",
                    "properties": {
                        "language": {"type": "string", "enum": ["shell"]},
                        "script": {"type": "string", "description": "Script content"},
                        "save_output": {"type": "boolean", "default": False},
                    },
                    "required": ["language", "script"],
                }
            })
            tools.append({
                "name": "computer_control",
                "description": "Control computer using shell commands",
                "inputSchema": {
                    "type": "object",
                    "properties": {
                        "script": {"type": "string", "description": "Shell script"},
                        "save_output": {"type": "boolean", "default": False},
                    },
                    "required": ["script"],
                }
            })

        tools.extend([
            {
                "name": "xlsx_tool",
                "description": "Process Excel files: list worksheets, read cells, find text, update cells",
                "inputSchema": {
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": "Path to XLSX file"},
                        "operation": {"type": "string", "enum": ["list_worksheets", "get_columns", "get_range", "find_text", "update_cell", "get_cell", "save"]},
                        "worksheet": {"type": "string", "description": "Worksheet name"},
                        "range": {"type": "string", "description": "Cell range (e.g., 'A1:C10')"},
                        "search_text": {"type": "string", "description": "Text to search"},
                        "row": {"type": "integer", "description": "Row number"},
                        "col": {"type": "integer", "description": "Column number"},
                        "value": {"type": "string", "description": "Value for update_cell"},
                    },
                    "required": ["path", "operation"],
                }
            },
            {
                "name": "docx_tool",
                "description": "Process Word documents: extract text, create/update documents",
                "inputSchema": {
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": "Path to DOCX file"},
                        "operation": {"type": "string", "enum": ["extract_text", "update_doc"]},
                        "content": {"type": "string", "description": "Content for update_doc"},
                    },
                    "required": ["path", "operation"],
                }
            },
            {
                "name": "pdf_tool",
                "description": "Process PDF files: extract text or images",
                "inputSchema": {
                    "type": "object",
                    "properties": {
                        "path": {"type": "string", "description": "Path to PDF file"},
                        "operation": {"type": "string", "enum": ["extract_text", "extract_images"]},
                    },
                    "required": ["path", "operation"],
                }
            },
        ])

        return tools

    async def call_tool(self, name: str, arguments: Dict[str, Any]) -> Dict[str, Any]:
        """Call a tool"""
        if not self._initialized:
            await self.initialize()

        handlers = {
            "web_scrape": self._web_scrape,
            "cache": self._cache,
            "automation_script": self._automation_script,
            "computer_control": self._computer_control,
            "xlsx_tool": self._xlsx_tool,
            "docx_tool": self._docx_tool,
            "pdf_tool": self._pdf_tool,
        }

        if name not in handlers:
            return {"error": f"Unknown tool: {name}"}

        try:
            return await handlers[name](arguments)
        except Exception as e:
            return {"error": str(e)}

    async def _web_scrape(self, args: Dict[str, Any]) -> Dict[str, Any]:
        """Fetch web content"""
        if not HTTPX_AVAILABLE or not self.http_client:
            return {"error": "httpx required for web scraping. Install with: pip install httpx"}

        url = args.get("url", "")
        save_as = args.get("save_as", "text")

        if not url:
            return {"error": "Missing 'url' parameter"}

        try:
            response = await self.http_client.get(url, headers={"Accept": "*/*"})
            response.raise_for_status()

            if save_as == "json":
                try:
                    content = response.json()
                    content_bytes = json.dumps(content, indent=2).encode("utf-8")
                    ext = "json"
                except Exception:
                    return {"error": "Response is not valid JSON"}

            elif save_as == "binary":
                content_bytes = response.content
                ext = "bin"
            else:
                content_bytes = response.text.encode("utf-8")
                ext = "txt"

            cache_path = self._get_cache_path("web", ext)
            cache_path.write_bytes(content_bytes)

            return {"content": [{"type": "text", "text": f"Content saved to: {cache_path}"}]}
        except Exception as e:
            return {"error": f"Failed to fetch URL: {e}"}

    async def _cache(self, args: Dict[str, Any]) -> Dict[str, Any]:
        """Manage cached files"""
        command = args.get("command", "")
        path = args.get("path", "")

        if command == "list":
            if not self.cache_dir.exists():
                return {"content": [{"type": "text", "text": "Cache is empty"}]}

            files = sorted(self.cache_dir.iterdir(), key=lambda x: x.stat().st_mtime, reverse=True)
            file_list = [f"{f.name} ({f.stat().st_size} bytes)" for f in files if f.is_file()]

            if not file_list:
                return {"content": [{"type": "text", "text": "Cache is empty"}]}

            return {"content": [{"type": "text", "text": "Cached files:\n" + "\n".join(file_list)}]}

        elif command == "view":
            if not path:
                return {"error": "Missing 'path' parameter"}

            path_obj = Path(path)
            if not path_obj.exists():
                return {"error": f"File not found: {path}"}

            try:
                content = path_obj.read_text(encoding="utf-8")
                return {"content": [{"type": "text", "text": f"Content of {path}:\n\n{content}"}]}
            except Exception:
                return {"content": [{"type": "text", "text": f"Binary file: {path}"}]}

        elif command == "delete":
            if not path:
                return {"error": "Missing 'path' parameter"}

            path_obj = Path(path)
            if path_obj.exists():
                path_obj.unlink()
            return {"content": [{"type": "text", "text": f"Deleted: {path}"}]}

        elif command == "clear":
            if self.cache_dir.exists():
                shutil.rmtree(self.cache_dir)
                self.cache_dir.mkdir(parents=True, exist_ok=True)
            return {"content": [{"type": "text", "text": "Cache cleared"}]}

        else:
            return {"error": f"Unknown command: {command}"}

    async def _automation_script(self, args: Dict[str, Any]) -> Dict[str, Any]:
        """Run automation script"""
        language = args.get("language", "shell")
        script = args.get("script", "")
        save_output = args.get("save_output", False)

        if not script:
            return {"error": "Missing 'script' parameter"}

        temp_dir = tempfile.mkdtemp()

        try:
            if language == "powershell":
                script_path = Path(temp_dir) / "script.ps1"
                script_path.write_text(script, encoding="utf-8")
                result = subprocess.run(
                    ["powershell", "-NoProfile", "-NonInteractive", "-File", str(script_path)],
                    capture_output=True,
                    text=True,
                    timeout=60,
                )
            elif language in ("shell", "batch"):
                ext = "bat" if language == "batch" and self._os == "nt" else "sh"
                script_path = Path(temp_dir) / f"script.{ext}"
                script_path.write_text(script, encoding="utf-8")
                if self._os != "nt":
                    os.chmod(script_path, 0o755)
                result = subprocess.run(
                    ["sh" if self._os != "nt" else "cmd", "/c", str(script_path)] if self._os == "nt" else ["sh", str(script_path)],
                    capture_output=True,
                    text=True,
                    timeout=60,
                )
            else:
                return {"error": f"Unsupported language: {language}"}

            output = result.stdout + result.stderr
            response = f"Script completed{' successfully' if result.returncode == 0 else ' with errors'}.\n\nOutput:\n{output}"

            if save_output and output.strip():
                cache_path = self._get_cache_path("script_output", "txt")
                cache_path.write_text(output, encoding="utf-8")
                response += f"\n\nOutput saved to: {cache_path}"

            return {"content": [{"type": "text", "text": response}]}

        except subprocess.TimeoutExpired:
            return {"error": "Script timed out after 60 seconds"}
        except Exception as e:
            return {"error": f"Script failed: {e}"}
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

    async def _computer_control(self, args: Dict[str, Any]) -> Dict[str, Any]:
        """Execute system control script"""
        script = args.get("script", "")
        save_output = args.get("save_output", False)

        if not script:
            return {"error": "Missing 'script' parameter"}

        try:
            if self._os == "nt":
                result = subprocess.run(
                    ["powershell", "-NoProfile", "-NonInteractive", "-Command", script],
                    capture_output=True,
                    text=True,
                    timeout=60,
                )
            else:
                result = subprocess.run(
                    ["sh", "-c", script],
                    capture_output=True,
                    text=True,
                    timeout=60,
                )

            output = result.stdout + result.stderr
            response = f"Script completed{' successfully' if result.returncode == 0 else ' with errors'}.\n\nOutput:\n{output}"

            if save_output and output.strip():
                cache_path = self._get_cache_path("automation_output", "txt")
                cache_path.write_text(output, encoding="utf-8")
                response += f"\n\nOutput saved to: {cache_path}"

            return {"content": [{"type": "text", "text": response}]}

        except subprocess.TimeoutExpired:
            return {"error": "Script timed out after 60 seconds"}
        except Exception as e:
            return {"error": f"Script failed: {e}"}

    async def _xlsx_tool(self, args: Dict[str, Any]) -> Dict[str, Any]:
        """Process Excel files"""
        path = args.get("path", "")
        operation = args.get("operation", "")

        if not path:
            return {"error": "Missing 'path' parameter"}

        try:
            import openpyxl
            wb = openpyxl.load_workbook(path)
        except Exception as e:
            return {"error": f"Failed to open Excel file: {e}"}

        try:
            if operation == "list_worksheets":
                sheets = [(name, ws.max_row, ws.max_column) for name, ws in wb.items()]
                result = [f"Worksheets in {path}:"]
                for name, rows, cols in sheets:
                    result.append(f"  - {name}: {rows} rows x {cols} columns")
                return {"content": [{"type": "text", "text": "\n".join(result)}]}

            elif operation == "get_columns":
                ws_name = args.get("worksheet")
                ws = wb[ws_name] if ws_name else wb.active
                columns = [str(cell.value) for cell in ws[1]]
                return {"content": [{"type": "text", "text": f"Columns: {columns}"}]}

            elif operation == "get_range":
                ws_name = args.get("worksheet")
                range_str = args.get("range", "A1:Z100")
                ws = wb[ws_name] if ws_name else wb.active
                data = []
                for row in ws[range_str]:
                    data.append([str(cell.value) for cell in row])
                return {"content": [{"type": "text", "text": json.dumps(data, indent=2)}]}

            elif operation == "find_text":
                ws_name = args.get("worksheet")
                search_text = args.get("search_text", "")
                ws = wb[ws_name] if ws_name else wb.active
                matches = []
                for row in ws.iter_rows():
                    for cell in row:
                        if cell.value and search_text.lower() in str(cell.value).lower():
                            matches.append(f"Row {cell.row}, Col {cell.column}: {cell.value}")
                if matches:
                    return {"content": [{"type": "text", "text": f"Found {len(matches)} matches:\n" + "\n".join(matches[:50])}]}
                return {"content": [{"type": "text", "text": "No matches found"}]}

            elif operation == "update_cell":
                ws_name = args.get("worksheet", "Sheet1")
                row = args.get("row", 1)
                col = args.get("col", 1)
                value = args.get("value", "")
                ws = wb[ws_name]
                ws.cell(row=row, column=col, value=value)
                wb.save(path)
                return {"content": [{"type": "text", "text": f"Updated cell ({row}, {col}) in {ws_name}"}]}

            elif operation == "get_cell":
                ws_name = args.get("worksheet")
                row = args.get("row", 1)
                col = args.get("col", 1)
                ws = wb[ws_name] if ws_name else wb.active
                cell = ws.cell(row=row, column=col)
                return {"content": [{"type": "text", "text": f"Value: {cell.value}, Formula: {cell.value if not cell.data_type == 'f' else cell.value}"}]}

            elif operation == "save":
                wb.save(path)
                return {"content": [{"type": "text", "text": "File saved successfully"}]}

            else:
                return {"error": f"Unknown operation: {operation}"}

        except Exception as e:
            return {"error": f"Operation failed: {e}"}

    async def _docx_tool(self, args: Dict[str, Any]) -> Dict[str, Any]:
        """Process Word documents"""
        path = args.get("path", "")
        operation = args.get("operation", "")

        if not path:
            return {"error": "Missing 'path' parameter"}

        try:
            from docx import Document
        except ImportError:
            return {"error": "python-docx required for DOCX processing. Install with: pip install python-docx"}

        try:
            if operation == "extract_text":
                if Path(path).exists():
                    doc = Document(path)
                    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                    return {"content": [{"type": "text", "text": f"Document: {path}\n\n" + "\n".join(paragraphs[:100])}]}
                else:
                    return {"content": [{"type": "text", "text": "New document will be created"}]}

            elif operation == "update_doc":
                content = args.get("content", "")
                doc = Document() if not Path(path).exists() else Document(path)
                for para in content.split("\n\n"):
                    doc.add_paragraph(para)
                doc.save(path)
                return {"content": [{"type": "text", "text": f"Document saved to {path}"}]}

            else:
                return {"error": f"Unknown operation: {operation}"}

        except Exception as e:
            return {"error": f"Operation failed: {e}"}

    async def _pdf_tool(self, args: Dict[str, Any]) -> Dict[str, Any]:
        """Process PDF files"""
        path = args.get("path", "")
        operation = args.get("operation", "")

        if not path:
            return {"error": "Missing 'path' parameter"}

        try:
            import PyPDF2
        except ImportError:
            return {"error": "PyPDF2 required for PDF processing. Install with: pip install PyPDF2"}

        try:
            if operation == "extract_text":
                with open(path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    text_parts = []
                    for page in reader.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                    text = "\n".join(text_parts)
                return {"content": [{"type": "text", "text": f"Extracted {len(text)} characters from PDF"}]}

            elif operation == "extract_images":
                return {"content": [{"type": "text", "text": "Image extraction requires additional libraries (pdf2image)"}]}

            else:
                return {"error": f"Unknown operation: {operation}"}

        except Exception as e:
            return {"error": f"Operation failed: {e}"}

    async def close(self) -> None:
        """Close extension"""
        self._initialized = False
        if self.http_client:
            await self.http_client.aclose()


def create_computer_controller_extension() -> ComputerControllerPlatformExtension:
    """Create Computer Controller Platform Extension"""
    return ComputerControllerPlatformExtension()
